import sys
import os
import base64
import zlib
import urllib.request

try:
    import docx
    from docx.shared import Inches, Pt
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx
    from docx.shared import Inches, Pt

def download_kroki_image(mermaid_code, filename):
    print(f"Downloading {filename}...")
    compressed = zlib.compress(mermaid_code.encode('utf-8'), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
    url = f"https://kroki.io/mermaid/png/{encoded}"
    image_path = os.path.join(os.getcwd(), filename)
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(image_path, 'wb') as out_file:
        out_file.write(response.read())
    return image_path

# ================= KROKI DIAGRAMS =================
diag1_arch = """
graph LR
    UI["Frontend<br>ReactJS + Tailwind"] <-->|"REST API JSON"| API["Backend<br>Node.js Express"]
    API <-->|"Prisma ORM"| DB[("Database<br>SQLite")]
"""
img_arch = download_kroki_image(diag1_arch, "img_arch.png")

diag2_test = """
sequenceDiagram
    actor Khach
    actor Admin
    Khach->>Frontend: Click 'Dat san'
    Frontend->>Backend: POST /api/bookings
    Backend->>Database: Kiem tra trung lich & Luu DB
    Database-->>Backend: Thanh cong
    Backend-->>Frontend: 201 Created
    Frontend->>Khach: Hien thi ma QR thanh toan
    Admin->>Backend: GET /api/bookings
    Backend-->>Admin: Cap nhat Danh sach don moi
"""
img_test = download_kroki_image(diag2_test, "img_test.png")

diag3_local = """
graph TD
    User(("Nguoi dung")) --> |"Click dup"| B["run_web.bat"]
    B --> |"Tien trinh 1"| C["Frontend ReactJS<br>Port 5173"]
    B --> |"Tien trinh 2"| D["Backend Node.js<br>Port 3000"]
    D --> |"Truy van"| E[("SQLite Database")]
    C <--> |"Goi API"| D
    B -.-> |"Cho 3s, tu dong mo"| F["Trinh duyet Web"]
    F --> |"Hien thi UI"| C
"""
img_local = download_kroki_image(diag3_local, "img_local.png")

diag4_prod = """
graph TD
    Client(("Khach hang<br>& Admin")) -->|"Truy cap HTTPS"| CDN["Frontend - Vercel/Netlify<br>Static React Files"]
    CDN -->|"Giao dien"| Client
    Client -->|"Goi API REST"| LB["Nginx Proxy / API Gateway"]
    LB --> Backend["Backend Server - Node.js<br>AWS EC2 / VPS"]
    Backend -->|"Doc/Ghi du lieu"| DB[("Production Database<br>PostgreSQL / MySQL")]
"""
img_prod = download_kroki_image(diag4_prod, "img_prod.png")

diag5_maint = """
graph TD
    Start(("Bat dau")) --> A["Giam sat he thong<br>(Monitoring)"]
    A --> B{"Phan loai Yeu cau"}
    B -->|"Moi truong doi"| C["Bao tri Thich ung<br>(Adaptive)"]
    B -->|"Nguy co tiem an"| D["Bao tri Phong ngua<br>(Preventive)"]
    C --> C1["Cap nhat Thu vien"]
    D --> D1["Refactor & Backup DB"]
    C1 --> E["Kiem thu & Trien khai"]
    D1 --> E
    E --> A
"""
img_maint = download_kroki_image(diag5_maint, "img_maint.png")

# ================= TẠO WORD =================
doc = docx.Document()

def add_heading(text, level): return doc.add_heading(text, level=level)
def add_paragraph(text, bold=False, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    return p

def insert_image(img_path, title):
    doc.add_picture(img_path, width=Inches(5.5))
    p = doc.add_paragraph(title)
    p.alignment = 1

# --- Chương 3 ---
add_heading('CHƯƠNG 3: CÀI ĐẶT VÀ KIỂM THỬ HỆ THỐNG', 1)
add_heading('3.1. Cài đặt hệ thống (System Implementation)', 2)
add_paragraph('Giai đoạn cài đặt hệ thống là bước chuyển giao mang tính quyết định, nơi các thiết kế kiến trúc và cơ sở dữ liệu được dịch sang ngôn ngữ lập trình. Quá trình này đòi hỏi sự phối hợp nhịp nhàng của các thành viên trong nhóm theo luồng công việc Agile/Scrum và khả năng làm chủ các nền tảng công nghệ hiện đại.')

add_heading('3.1.1. Lựa chọn công nghệ và Thiết lập môi trường phát triển', 3)
add_paragraph('Hệ thống lựa chọn kiến trúc ứng dụng trang đơn (Single Page Application - SPA) chạy theo mô hình Client-Server hiện đại. Dưới đây là chi tiết công nghệ cốt lõi được áp dụng:')
add_paragraph('Frontend (Giao diện): Sử dụng ReactJS kết hợp TypeScript. Giao diện được thiết kế bằng Tailwind CSS.', style='List Bullet')
add_paragraph('Backend (Máy chủ xử lý): Phát triển trên Node.js kết hợp framework Express.js.', style='List Bullet')
add_paragraph('Cơ sở dữ liệu (Database): Sử dụng SQLite kết hợp với công cụ Prisma ORM.', style='List Bullet')

add_heading('3.1.2. Tổ chức mã nguồn và Tiêu chuẩn hợp tác (Coding Conventions)', 3)
add_paragraph('Quy chuẩn được áp dụng nghiêm ngặt: Frontend chia thành Component (PascalCase), Backend tổ chức theo route (camelCase), sử dụng Git nhánh riêng rẽ.')

add_heading('3.1.3. Hiện thực hóa các Module nghiệp vụ trọng yếu', 3)
add_paragraph('Hệ thống sử dụng kiến trúc CSDL thực tế thay vì bộ nhớ tạm thời:')
insert_image(img_arch, 'Hình 3.1: Sơ đồ Kiến trúc luồng dữ liệu ReactJS - Node.js - SQLite')

add_heading('3.2. Kiểm thử phần mềm (Software Testing)', 2)
add_heading('3.2.1. Kiểm thử tích hợp hệ thống (Integration Testing)', 3)
add_paragraph('Kiểm thử tích hợp tập trung kiểm tra luồng truyền dữ liệu xuyên suốt giữa các Module:')
insert_image(img_test, 'Hình 3.2: Sơ đồ Tuần tự (Sequence Diagram) luồng Đặt sân và Kiểm thử Tích hợp')

# --- Chương 4 ---
add_heading('CHƯƠNG 4: VẬN HÀNH VÀ BẢO TRÌ HỆ THỐNG', 1)
add_heading('4.1. Quy trình Triển khai và Vận hành', 2)

add_heading('4.1.1. Vận hành trong môi trường nội bộ (Local Automation)', 3)
add_paragraph('Hệ thống thiết lập script tự động chạy đồng thời Backend và Frontend:')
insert_image(img_local, 'Hình 4.1: Sơ đồ luồng khởi chạy tự động Local Automation')

add_heading('4.1.2. Vận hành trong môi trường thực tế (Production)', 3)
add_paragraph('Kiến trúc được đưa lên Cloud với Nginx, Docker và PostgreSQL để chịu tải:')
insert_image(img_prod, 'Hình 4.2: Sơ đồ Triển khai hạ tầng Môi trường Thực tế (Production)')

add_heading('4.2. Chiến lược Bảo trì phần mềm', 2)
add_heading('4.2.1. Bảo trì khắc phục (Corrective Maintenance)', 3)
add_paragraph('Sửa chữa các lỗi logic hoặc lỗi hiển thị phát sinh khi người dùng thao tác thực tế.')

add_heading('4.2.2. Bảo trì Thích ứng và Phòng ngừa', 3)
add_paragraph('Bao gồm cập nhật thư viện phòng ngừa lỗ hổng bảo mật, và tái cấu trúc mã nguồn.')
insert_image(img_maint, 'Hình 4.3: Vòng đời và Quy trình Bảo trì Hệ thống')

add_heading('4.3. Quản trị rủi ro và Đảm bảo an toàn dữ liệu', 2)
add_paragraph('Mật khẩu được băm (bcrypt) và API được bảo vệ bằng Token. Cài đặt Rate Limit chống spam.')

# --- Chương 5 ---
add_heading('CHƯƠNG 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN', 1)
add_heading('5.1. Đánh giá kết quả đạt được', 2)
add_paragraph('Hoàn thành hệ thống đúng tiến độ, đáp ứng 100% nghiệp vụ và áp dụng tốt công nghệ MERN/PERN stack.')

add_heading('5.2. Những mặt hạn chế còn tồn tại', 2)
add_paragraph('Tính năng thanh toán chỉ hiển thị mã VietQR tĩnh, chưa gọi API trực tiếp đến cổng thanh toán.')

add_heading('5.3. Định hướng phát triển', 2)
add_paragraph('Phát triển Mobile App với React Native và tích hợp AI dự báo giờ vàng.')

filepath = os.path.join(os.getcwd(), 'Bao_Cao_Chuong_3_4_5_Kem_Hinh_Anh.docx')
doc.save(filepath)
print(f"File created at {filepath}")
