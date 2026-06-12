import sys
import subprocess
import os
import base64
import zlib
import urllib.request

try:
    import docx
    from docx.shared import Inches
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx
    from docx.shared import Inches

# 1. Download Diagram from Kroki
mermaid_code = """
graph TD
    Start(("Bat dau")) --> A["Giam sat he thong<br>(Monitoring)"]
    A --> B{"Phan loai Yeu cau"}
    B -->|"Moi truong thay doi"| C["Bao tri Thich ung<br>(Adaptive)"]
    B -->|"Nguy co tiem an"| D["Bao tri Phong ngua<br>(Preventive)"]
    C --> C1["Cap nhat Thu vien<br>React/Node.js"]
    C --> C2["Toi uu Responsive<br>Cac thiet bi moi"]
    D --> D1["Refactor Code<br>voi TypeScript"]
    D --> D2["Backup Database<br>Dinh ky"]
    C1 --> E["Kiem thu (Testing)<br>& Trien khai"]
    C2 --> E
    D1 --> E
    D2 --> E
    E --> A
"""
print("Downloading maintenance diagram...")
compressed = zlib.compress(mermaid_code.encode('utf-8'), 9)
encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
url = f"https://kroki.io/mermaid/png/{encoded}"

image_path = os.path.join(os.getcwd(), 'diagram_maintenance.png')
req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
with urllib.request.urlopen(req) as response, open(image_path, 'wb') as out_file:
    out_file.write(response.read())

# 2. Update Word Document
doc = docx.Document()

# Tiêu đề chính
doc.add_heading('Bảo trì thích ứng và Phòng ngừa (Adaptive & Preventive Maintenance)', level=2)

doc.add_paragraph('Bảo trì phần mềm không chỉ đơn thuần là sửa lỗi (Fix bugs) khi hệ thống đã hỏng, mà còn là việc chủ động nâng cấp và phòng ngừa rủi ro để đảm bảo tính sẵn sàng cao (High Availability). Đối với Hệ thống quản lý sân cầu lông, chiến lược bảo trì được chia làm hai mảng chính:')

# Chèn Hình ảnh
doc.add_picture(image_path, width=Inches(5.5))
p_img = doc.add_paragraph('Hình: Vòng đời và Quy trình Bảo trì hệ thống')
p_img.alignment = 1 # Center align

# 1. Bảo trì Thích ứng
doc.add_heading('1. Bảo trì Thích ứng (Adaptive Maintenance)', level=3)
doc.add_paragraph('Bảo trì thích ứng tập trung vào việc cập nhật phần mềm để nó vẫn hoạt động tốt khi môi trường bên ngoài (hệ điều hành, phần cứng, trình duyệt web) thay đổi.')

p_ad1 = doc.add_paragraph('', style='List Bullet')
p_ad1.add_run('Cập nhật nền tảng và Thư viện: ').bold = True
p_ad1.add_run('Hệ thống được phát triển bằng ReactJS và Node.js. Mỗi năm, các trình duyệt web (Chrome, Safari) đều ra mắt phiên bản mới với các tiêu chuẩn bảo mật khắt khe hơn. Bảo trì thích ứng bao gồm việc nâng cấp thư viện (như Express, Prisma, Tailwind CSS) lên phiên bản mới nhất để không bị trình duyệt đánh dấu là "Không an toàn" và loại bỏ các lỗ hổng bảo mật (Vulnerabilities).')

p_ad2 = doc.add_paragraph('', style='List Bullet')
p_ad2.add_run('Tương thích Thiết bị mới (Responsive Adaptation): ').bold = True
p_ad2.add_run('Khi có các dòng điện thoại thông minh mới (màn hình gập, tablet với kích thước dị biệt) xuất hiện trên thị trường, giao diện UI/UX của hệ thống cần được căn chỉnh lại (bằng CSS Flexbox/Grid) để thích ứng mượt mà mà không làm vỡ bố cục hiển thị lịch đặt sân.')

# 2. Bảo trì Phòng ngừa
doc.add_heading('2. Bảo trì Phòng ngừa (Preventive Maintenance)', level=3)
doc.add_paragraph('Bảo trì phòng ngừa là các hoạt động chủ động can thiệp vào mã nguồn hoặc hạ tầng trước khi lỗi thực sự xảy ra, giúp tăng tuổi thọ và độ ổn định của hệ thống.')

p_pv1 = doc.add_paragraph('', style='List Bullet')
p_pv1.add_run('Tái cấu trúc mã nguồn (Refactoring) bằng TypeScript: ').bold = True
p_pv1.add_run('Mã nguồn liên tục được dọn dẹp (Clean Code) và cấu trúc lại các Component. Việc sử dụng chặt chẽ TypeScript giúp phát hiện sớm các rủi ro về kiểu dữ liệu (Data Type) ngay từ lúc viết code (Compile time) thay vì để nó bùng phát lúc phần mềm đang chạy (Runtime).')

p_pv2 = doc.add_paragraph('', style='List Bullet')
p_pv2.add_run('Sao lưu Dữ liệu Tự động (Automated Backups): ').bold = True
p_pv2.add_run('Phòng ngừa rủi ro hỏng hóc đĩa cứng hoặc bị tấn công tống tiền (Ransomware) bằng cách lên lịch (Cron Job) sao lưu cơ sở dữ liệu định kỳ mỗi 12 giờ lên các dịch vụ lưu trữ đám mây an toàn (như AWS S3 hoặc Google Drive).')

p_pv3 = doc.add_paragraph('', style='List Bullet')
p_pv3.add_run('Giám sát hiệu năng (Performance Monitoring): ').bold = True
p_pv3.add_run('Cài đặt các công cụ theo dõi log tự động. Nếu phát hiện một tài khoản có tần suất đặt sân bất thường (dấu hiệu của tool spam/bot), hệ thống sẽ phòng ngừa bằng cách tạm khoá tài khoản và cảnh báo cho quản trị viên.')

filepath = os.path.join(os.getcwd(), 'Phan_Bao_Tri_Thich_Ung_Phong_Ngua.docx')
doc.save(filepath)
print(f"Da tao thanh cong file: {filepath}")
