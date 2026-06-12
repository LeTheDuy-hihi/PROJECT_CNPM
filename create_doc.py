import sys
import subprocess
import os

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx

doc = docx.Document()

# Tiêu đề
doc.add_heading('4.1.1. Vận hành trong môi trường nội bộ và Tự động hóa (Local Automation)', level=3)

# Đoạn a
p = doc.add_paragraph()
r = p.add_run('a) Kiến trúc triển khai nội bộ (Local Environment)')
r.bold = True

doc.add_paragraph('Hệ thống quản lý sân cầu lông được thiết kế để có thể hoạt động độc lập và toàn vẹn ngay trên máy tính cá nhân (localhost) của người dùng hoặc người quản trị mà không phụ thuộc vào máy chủ đám mây (Cloud) hay kết nối mạng Internet bên ngoài. Kiến trúc này được chia làm 3 thành phần chính chạy song song trên các cổng mạng (port) nội bộ:')

doc.add_paragraph('Máy chủ Giao diện (Frontend Server): Chạy trên nền tảng Node.js thông qua công cụ Vite (Port 5173), chịu trách nhiệm render giao diện người dùng (Customer Layout) và giao diện quản trị (Admin Dashboard).', style='List Bullet')
doc.add_paragraph('Máy chủ Xử lý (Backend API): Chạy trên cổng 3000, xử lý các nghiệp vụ logic (đặt sân, kiểm tra trùng lịch, xác thực đăng nhập) bằng Express.js.', style='List Bullet')
doc.add_paragraph('Cơ sở dữ liệu cục bộ (Local Database): Sử dụng hệ quản trị CSDL SQLite (tệp dev.db), được tích hợp trực tiếp thông qua ORM Prisma. Việc sử dụng SQLite giúp hệ thống lưu trữ dữ liệu bền vững ngay trong thư mục mã nguồn mà không yêu cầu người dùng phải cài đặt thêm các hệ quản trị phức tạp như MySQL hay SQL Server.', style='List Bullet')

# Đoạn b
p2 = doc.add_paragraph()
r2 = p2.add_run('b) Cơ chế Tự động hóa quá trình khởi chạy (Local Automation)')
r2.bold = True

doc.add_paragraph('Để giải quyết bài toán "khởi động nhiều thành phần cùng lúc" vốn rất phức tạp đối với người dùng cuối, hệ thống đã ứng dụng cơ chế Tự động hóa (Automation) thông qua hai kịch bản khởi chạy: Script Batch (run_web.bat) và Script Python (main.py). Cơ chế này tuân theo sơ đồ luồng "One-click Run" (Chạy với một cú click chuột):')

doc.add_paragraph('Phân luồng tiến trình (Multi-processing): Khi người dùng kích hoạt tệp main.py hoặc run_web.bat, hệ thống sẽ gọi các thư viện hệ thống (subprocess trong Python hoặc lệnh start trong CMD) để rẽ nhánh ra hai tiến trình (process) hoàn toàn độc lập.', style='List Number')
doc.add_paragraph('Khởi động đồng thời: Tiến trình 1 tự động điều hướng vào thư mục backend và kích hoạt lệnh npm run dev. Tiến trình 2 tự động điều hướng vào thư mục frontend và kích hoạt lệnh npm run dev.', style='List Number')
doc.add_paragraph('Cơ chế chờ đồng bộ (Delay & Sync): Hệ thống được lập trình để tự động "ngủ" (delay) trong khoảng 3 giây. Quá trình này đảm bảo cả hai máy chủ Backend và Frontend đều đã được khởi động hoàn tất trước khi chuyển sang bước tiếp theo.', style='List Number')
doc.add_paragraph('Tự động gọi Trình duyệt: Script sử dụng module webbrowser hoặc lệnh hệ thống để tự động đánh thức trình duyệt mặc định và truy cập thẳng vào địa chỉ ứng dụng: http://localhost:5173.', style='List Number')

# Đoạn c
p3 = doc.add_paragraph()
r3 = p3.add_run('c) Ưu điểm của mô hình')
r3.bold = True

doc.add_paragraph('Việc tích hợp cơ chế Local Automation trực tiếp vào dự án mang lại nhiều giá trị cốt lõi:')
doc.add_paragraph('Tính thân thiện (User-friendly): Bất kỳ ai đều có thể dễ dàng bật toàn bộ hệ thống lên chỉ bằng một thao tác click đúp chuột vào file run_web.bat mà không cần biết các lệnh lập trình rườm rà.', style='List Bullet')
doc.add_paragraph('Tính đóng gói cao: Toàn bộ CSDL, Backend và Frontend được gói gọn. Môi trường phát triển và môi trường thực thi (runtime) được hợp nhất một cách liền mạch.', style='List Bullet')

filepath = os.path.join(os.getcwd(), 'Phan_4_1_1.docx')
doc.save(filepath)
print(f"Đã tạo thành công file: {filepath}")
