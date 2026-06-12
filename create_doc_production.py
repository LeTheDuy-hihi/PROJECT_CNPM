import sys
import subprocess
import os
import base64
import zlib
import urllib.request

try:
    import docx
    from docx.shared import Inches
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx
    from docx.shared import Inches

# 1. Download Diagram from Kroki
mermaid_code = """
graph TD
    Client(("Khach hang<br>& Admin")) -->|"Truy cap qua HTTPS"| CDN["Frontend - Vercel/Netlify<br>Static React Files"]
    CDN -->|"Tra ve Giao dien"| Client
    Client -->|"Goi API REST/JSON"| LB["Nginx Proxy / API Gateway"]
    LB --> Backend["Backend Server - Node.js<br>AWS EC2 / VPS"]
    Backend -->|"Doc/Ghi du lieu"| DB[("Production Database<br>PostgreSQL / MySQL")]
    Backend -.->|"Tich hop"| ThirdParty["Dich vu ngoai<br>(Thanh toan, Email)"]
"""
print("Downloading architecture diagram...")
compressed = zlib.compress(mermaid_code.encode('utf-8'), 9)
encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
url = f"https://kroki.io/mermaid/png/{encoded}"

image_path = os.path.join(os.getcwd(), 'diagram_production.png')
req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
with urllib.request.urlopen(req) as response, open(image_path, 'wb') as out_file:
    out_file.write(response.read())

# 2. Update Word Document
doc = docx.Document()

# Tiêu đề
doc.add_heading('4.1.2. Vận hành trong môi trường thực tế (Production Environment)', level=3)

# Đoạn giới thiệu
doc.add_paragraph('Khi triển khai hệ thống ra môi trường Internet để phục vụ khách hàng thực tế (Production), hệ thống cần được tái cấu trúc để đảm bảo tính bảo mật, hiệu năng cao và khả năng chịu tải (Scalability). Thay vì chạy cục bộ trên localhost, các thành phần sẽ được phân tán lên hạ tầng Điện toán đám mây (Cloud Computing).')

# Chèn Hình ảnh
doc.add_picture(image_path, width=Inches(6.0))
p_img = doc.add_paragraph('Hình 4.2: Sơ đồ kiến trúc triển khai hệ thống trên môi trường Thực tế (Production)')
p_img.alignment = 1 # Center align

# Cấu trúc
p = doc.add_paragraph()
r = p.add_run('Các thay đổi cốt lõi trong kiến trúc Production:')
r.bold = True

# Frontend
p_fe = doc.add_paragraph('', style='List Bullet')
p_fe.add_run('Máy chủ Giao diện (Frontend - ReactJS): ').bold = True
p_fe.add_run('Mã nguồn Frontend được biên dịch (build) thành các tệp tin tĩnh (HTML, CSS, JS tối ưu hoá). Các tệp này sẽ được đưa lên các nền tảng phân phối nội dung toàn cầu (CDN) như Vercel hoặc Netlify. Điều này giúp giao diện web tải với tốc độ chưa tới 1 giây bất kể vị trí địa lý của khách hàng.')

# Backend
p_be = doc.add_paragraph('', style='List Bullet')
p_be.add_run('Máy chủ Xử lý (Backend - Node.js): ').bold = True
p_be.add_run('API Server được đóng gói bằng Docker và triển khai lên các máy chủ ảo (VPS) như AWS EC2, DigitalOcean hoặc Render. Máy chủ được bảo vệ bởi một Nginx Reverse Proxy đóng vai trò điều hướng lưu lượng truy cập, cân bằng tải và chặn các cuộc tấn công mạng (DDoS).')

# Database
p_db = doc.add_paragraph('', style='List Bullet')
p_db.add_run('Cơ sở dữ liệu (Production Database): ').bold = True
p_db.add_run('Hệ quản trị CSDL SQLite (vốn chỉ dùng cho môi trường giả lập cục bộ) sẽ được nâng cấp lên PostgreSQL hoặc MySQL. Nhờ kiến trúc linh hoạt của Prisma ORM, việc chuyển đổi CSDL chỉ tốn một dòng cấu hình. PostgreSQL đảm bảo tính toàn vẹn dữ liệu khi có hàng trăm khách hàng cùng lúc vào đặt một khung giờ sân (giải quyết triệt để lỗi Concurrency).')

# Bảo mật
p_sec = doc.add_paragraph('', style='List Bullet')
p_sec.add_run('Tên miền và Bảo mật (Domain & SSL): ').bold = True
p_sec.add_run('Hệ thống được gắn liền với một tên miền chính thức (ví dụ: sancaulong.vn). Toàn bộ dữ liệu trao đổi giữa khách hàng và máy chủ được mã hóa thông qua giao thức HTTPS (chứng chỉ SSL/TLS), đảm bảo thông tin cá nhân và mật khẩu của người dùng không bị đánh cắp.')

filepath = os.path.join(os.getcwd(), 'Phan_4_1_2_Production.docx')
doc.save(filepath)
print(f"Đã tạo thành công file: {filepath}")
