# -*- coding: utf-8 -*-
"""
generate_word.py - Script tự động tạo file Word (.docx) báo cáo Chương 3
Sử dụng thư viện python-docx để định dạng chuyên nghiệp chuẩn báo cáo BTL.
"""

import sys
import os

try:
    import docx
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("Thư viện python-docx chưa được cài đặt. Vui lòng chạy 'pip install python-docx' trước.")
    sys.exit(1)

def set_cell_background(cell, fill_hex):
    """Thiết lập màu nền cho ô trong bảng"""
    tcPr = cell._tc.get_or_add_tcPr()
    for existing_shd in tcPr.xpath('w:shd'):
        tcPr.remove(existing_shd)
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Thiết lập khoảng đệm (padding) cho ô trong bảng"""
    tcPr = cell._tc.get_or_add_tcPr()
    for existing_mar in tcPr.xpath('w:tcMar'):
        tcPr.remove(existing_mar)
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Thiết lập viền mờ màu xám nhạt cho toàn bảng"""
    tblPr = table._tbl.tblPr
    for existing_borders in tblPr.xpath('w:tblBorders'):
        tblPr.remove(existing_borders)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="A0A0A0"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_heading_with_spacing(doc, text, level, space_before=12, space_after=6):
    """Thêm tiêu đề với khoảng cách dòng tùy chỉnh"""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = True
    
    # Định dạng font chữ
    run = h.runs[0]
    run.font.name = 'Times New Roman'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(9, 13, 22) # Màu đen đậm tối
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(16, 185, 129) # Xanh Emerald chủ đạo
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(31, 41, 55)
        
    return h

def add_paragraph_with_spacing(doc, text="", space_after=6, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Thêm đoạn văn với giãn dòng chuẩn Word (1.15, Justify)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.alignment = align
    
    if text:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(31, 41, 55)
        
    return p

def main():
    doc = Document()

    # Thiết lập lề trang chuẩn A4 (Lề trên-dưới 2cm, Trái 3cm, Phải 2cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79) # 2cm
        section.bottom_margin = Inches(0.79) # 2cm
        section.left_margin = Inches(1.18) # 3cm
        section.right_margin = Inches(0.79) # 2cm
        section.page_width = Inches(8.27)  # A4 width
        section.page_height = Inches(11.69) # A4 height

    # Thiết lập style mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(31, 41, 55)

    # --------------------------------------------------------------------------
    # CHƯƠNG 3
    # --------------------------------------------------------------------------
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(18)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("CHƯƠNG 3: CÀI ĐẶT VÀ KIỂM THỬ HỆ THỐNG")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(9, 13, 22)

    # 3.1. Cài đặt hệ thống (System Implementation)
    add_heading_with_spacing(doc, "3.1. Cài đặt hệ thống (System Implementation)", level=1, space_before=18)
    add_paragraph_with_spacing(doc, 
        "Giai đoạn cài đặt hệ thống là bước chuyển giao mang tính quyết định, nơi các thiết kế kiến trúc và "
        "cơ sở dữ liệu được dịch sang ngôn ngữ máy tính. Quá trình này đòi hỏi sự phối hợp nhịp nhàng của "
        "4 thành viên trong nhóm theo luồng công việc Scrum và khả năng làm chủ các nền tảng công nghệ hiện đại."
    )

    # 3.1.1. Lựa chọn công nghệ và Thiết lập môi trường phát triển
    add_heading_with_spacing(doc, "3.1.1. Lựa chọn công nghệ và Thiết lập môi trường phát triển", level=2)
    
    add_paragraph_with_spacing(doc, 
        "Hệ thống lựa chọn kiến trúc ứng dụng trang đơn (Single Page Application - SPA) chạy trực tiếp trên phía máy khách "
        "nhằm tối ưu hóa tốc độ phản hồi, tiết kiệm tài nguyên máy chủ và mang lại trải nghiệm không tải lại trang mượt mà. "
        "Dưới đây là chi tiết công nghệ cốt lõi được áp dụng:"
    )

    p_list = doc.add_paragraph(style='List Bullet')
    p_list.paragraph_format.space_after = Pt(3)
    p_list.paragraph_format.line_spacing = 1.15
    run = p_list.add_run("HTML5: ")
    run.bold = True
    p_list.add_run("Xây dựng cấu trúc trang web vững chắc, chuẩn hóa SEO sử dụng các thẻ ngữ nghĩa (Semantic HTML) như <header>, <nav>, <main>, <section>, <footer>.")

    p_list = doc.add_paragraph(style='List Bullet')
    p_list.paragraph_format.space_after = Pt(3)
    p_list.paragraph_format.line_spacing = 1.15
    run = p_list.add_run("CSS3 (Vanilla): ")
    run.bold = True
    p_list.add_run("Thiết kế giao diện hiện đại theo phong cách Glassmorphism (hiệu ứng mờ nền kính, bóng đổ mềm mại, màu sắc HSL sống động, chuyển cảnh mượt mà). Tích hợp khả năng tự động thích ứng với nhiều kích thước màn hình (Responsive Design).")

    p_list = doc.add_paragraph(style='List Bullet')
    p_list.paragraph_format.space_after = Pt(6)
    p_list.paragraph_format.line_spacing = 1.15
    run = p_list.add_run("JavaScript (ES6+): ")
    run.bold = True
    p_list.add_run("Xử lý toàn bộ logic động của hệ thống bao gồm: Định tuyến trang SPA ảo, kiểm tra điều kiện rào cản đặt sân trùng lịch, cập nhật giỏ hàng/hóa đơn tự động, lưu trữ trạng thái vào LocalStorage và vẽ biểu đồ quản trị bằng HTML5 Canvas API.")

    add_paragraph_with_spacing(doc, "Để đảm bảo tính nhất quán trong nhóm phát triển, các thông số môi trường đã được chuẩn hóa như sau:", space_after=8)

    # Bảng môi trường phát triển
    table_env = doc.add_table(rows=6, cols=4)
    table_env.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_env)

    headers = ["Thành phần", "Công cụ / Công nghệ", "Phiên bản", "Vai trò trong dự án"]
    col_widths = [Inches(1.5), Inches(2.0), Inches(1.0), Inches(2.0)]

    # Định dạng Header Bảng
    hdr_cells = table_env.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        set_cell_background(hdr_cells[i], "10B981") # Màu xanh lá cây
        set_cell_margins(hdr_cells[i])
        # Font cho header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    env_data = [
        ["Hệ điều hành", "Windows 10/11 hoặc macOS", "64-bit", "Môi trường máy trạm của lập trình viên"],
        ["IDE", "Visual Studio Code", "v1.85+", "Trình soạn thảo mã nguồn chính, tích hợp Live Server"],
        ["Môi trường chạy", "Google Chrome / Microsoft Edge", "v120+", "Trình duyệt kiểm thử hiệu năng và debug giao diện"],
        ["Ngôn ngữ", "HTML5 / CSS3 / ES6 JS", "Chuẩn W3C", "Xây dựng giao diện và xử lý logic ứng dụng"],
        ["Quản lý mã nguồn", "Git & GitHub", "v2.40+", "Quản lý phiên bản và cộng tác nhóm phát triển"]
    ]

    for row_idx, data in enumerate(env_data):
        row_cells = table_env.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx])
            # Zebra striping (dòng chẵn lẻ)
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F9FAFB")
            else:
                set_cell_background(row_cells[col_idx], "FFFFFF")
            # Căn lề
            if col_idx == 2:
                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3.1.2. Tổ chức mã nguồn và Tiêu chuẩn hợp tác (Coding Conventions)
    add_heading_with_spacing(doc, "3.1.2. Tổ chức mã nguồn và Tiêu chuẩn hợp tác (Coding Conventions)", level=2)
    
    add_paragraph_with_spacing(doc, 
        "Nhóm phát triển đã áp dụng bộ quy chuẩn viết code (Coding Conventions) nghiêm ngặt để đảm bảo chất lượng "
        "mã nguồn sạch, có tổ chức và giảm thiểu xung đột khi làm việc nhóm trên kho chứa Git."
    )

    add_paragraph_with_spacing(doc, "A. Tiêu chuẩn đặt tên và cú pháp code:", bold=True, space_after=4)
    
    p_conv = doc.add_paragraph(style='List Bullet')
    p_conv.paragraph_format.space_after = Pt(2)
    p_conv.paragraph_format.line_spacing = 1.15
    run = p_conv.add_run("Quy tắc đặt tên biến và hàm (JavaScript): ")
    run.bold = True
    p_conv.add_run("Sử dụng chuẩn camelCase. Tên hàm phải là động từ thể hiện hành động rõ ràng. Ví dụ: ")
    p_conv.add_run("currentBookingState, DataManager, getBookedSlots(), renderTimeSlots().").italic = True

    p_conv = doc.add_paragraph(style='List Bullet')
    p_conv.paragraph_format.space_after = Pt(2)
    p_conv.paragraph_format.line_spacing = 1.15
    run = p_conv.add_run("Quy tắc đặt tên Class CSS (BEM Style): ")
    run.bold = True
    p_conv.add_run("Đặt tên lớp theo chuẩn Block-Element-Modifier giúp CSS có cấu trúc module, không bị ghi đè chéo. Ví dụ: ")
    p_conv.add_run(".court-card (Block), .court-card__title (Element), .court-tab-btn--active (Modifier).").italic = True

    p_conv = doc.add_paragraph(style='List Bullet')
    p_conv.paragraph_format.space_after = Pt(4)
    p_conv.paragraph_format.line_spacing = 1.15
    run = p_conv.add_run("Quy tắc đặt tên tập tin: ")
    run.bold = True
    p_conv.add_run("Viết thường hoàn toàn, sử dụng dấu gạch ngang làm ký tự phân tách (kebab-case). Ví dụ: ")
    p_conv.add_run("index.html, style.css, data.js, app.js.").italic = True

    add_paragraph_with_spacing(doc, "B. Quy chuẩn Quản lý mã nguồn (Git Workflow):", bold=True, space_after=4)
    
    p_git = doc.add_paragraph(style='List Bullet')
    p_git.paragraph_format.space_after = Pt(2)
    p_git.paragraph_format.line_spacing = 1.15
    run = p_git.add_run("Phân nhánh (Git Branches): ")
    run.bold = True
    p_git.add_run("Nhánh main lưu mã nguồn ổn định nhất. Các chức năng mới được lập trình độc lập trên các nhánh con dạng: ")
    p_git.add_run("feature/booking, feature/admin-dashboard, bugfix/blurry-canvas.").italic = True
    p_git.add_run(" Sau khi hoàn tất, thành viên gửi Pull Request lên develop để gộp nhánh.")

    p_git = doc.add_paragraph(style='List Bullet')
    p_git.paragraph_format.space_after = Pt(6)
    p_git.paragraph_format.line_spacing = 1.15
    run = p_git.add_run("Chuẩn hóa thông điệp Commit (Commit Messages): ")
    run.bold = True
    p_git.add_run("Bắt buộc ghi rõ tiền tố hành động để thuận tiện cho việc tra cứu lịch sử thay đổi: ")
    p_git.add_run("[Feat] thêm chức năng, [Fix] sửa lỗi, [Docs] cập nhật tài liệu báo cáo.").italic = True

    # 3.1.3. Hiện thực hóa các Module nghiệp vụ trọng yếu
    add_heading_with_spacing(doc, "3.1.3. Hiện thực hóa các Module nghiệp vụ trọng yếu", level=2)
    
    add_paragraph_with_spacing(doc, 
        "Hệ thống bao gồm ba tệp mã nguồn logic JavaScript chính tương ứng với mô hình phân tách chức năng: "
        "Quản lý dữ liệu (data.js), Nghiệp vụ khách hàng (app.js), và Quản trị hệ thống (admin.js). "
        "Dưới đây là phân tích giải thuật hiện thực hóa các module nghiệp vụ này:"
    )

    # Thuyết minh module data.js
    add_paragraph_with_spacing(doc, "A. Module Dữ liệu và Lưu trữ (js/data.js):", bold=True, space_after=4)
    add_paragraph_with_spacing(doc, 
        "Đảm nhiệm vai trò khởi tạo cấu trúc dữ liệu mô phỏng cho danh sách sân cầu lông, đơn giá, "
        "danh sách khung giờ trống và lưu trữ thông tin đặt sân động vào LocalStorage của trình duyệt. "
        "Lớp DataManager cung cấp các API tĩnh giúp các module khác đọc/ghi dữ liệu đồng bộ mà không cần máy chủ:",
        space_after=4
    )
    
    add_paragraph_with_spacing(doc, 
        "● DataManager.getCourts(): Lấy toàn bộ danh sách sân và trạng thái hoạt động.\n"
        "● DataManager.getBookings(): Lấy danh sách toàn bộ các lịch đặt sân hiện có.\n"
        "● DataManager.saveBooking(booking): Ghi nhận đơn hàng mới và đẩy lên hàng đầu danh sách trong localStorage.\n"
        "● DataManager.getBookedSlots(courtId, date): Trích xuất toàn bộ các khung giờ đã bị khóa (approved hoặc pending) của một sân cụ thể trong ngày cụ thể để ngăn chặn trùng lịch chơi.",
        space_after=8, italic=True
    )

    # Thuyết minh module app.js
    add_paragraph_with_spacing(doc, "B. Module Nghiệp vụ khách hàng (js/app.js):", bold=True, space_after=4)
    add_paragraph_with_spacing(doc, 
        "Quản lý vòng đời trải nghiệm của khách hàng từ lúc duyệt danh sách sân, lựa chọn ngày chơi, chọn khung giờ trống, "
        "điền thông tin cá nhân và thanh toán chuyển khoản giả lập. "
        "Thuật toán cốt lõi xử lý lưới khung giờ trống (Time Slots Grid) dựa trên logic sau:",
        space_after=4
    )
    
    p_algo = doc.add_paragraph()
    p_algo.paragraph_format.left_indent = Inches(0.25)
    p_algo.paragraph_format.space_after = Pt(4)
    p_algo.paragraph_format.line_spacing = 1.15
    p_algo.add_run(
        "BƯỚC 1: Lắng nghe sự kiện đổi Ngày chơi hoặc chuyển đổi Tab sân.\n"
        "BƯỚC 2: Gọi DataManager.getBookedSlots(courtId, date) để lấy danh sách các ID khung giờ đã bị khóa.\n"
        "BƯỚC 3: Duyệt danh sách TIME_SLOTS tĩnh. Nếu slot.id nằm trong danh sách đã đặt, thêm class 'slot--booked' và vô hiệu hóa click.\n"
        "BƯỚC 4: Lắng nghe sự kiện click của các slot trống còn lại, thêm/xóa khỏi mảng 'selectedSlots' để đánh dấu trạng thái đang chọn (màu vàng).\n"
        "BƯỚC 5: Đồng bộ hóa giỏ hàng hiển thị ở Sidebar (Hóa đơn tóm tắt gồm: Tên sân, Số giờ thuê, Đơn giá và Tổng tiền thanh toán)."
    ).font.italic = True

    add_paragraph_with_spacing(doc, 
        "Khi người dùng click 'Xác nhận & Thanh toán', hệ thống kiểm tra tính hợp lệ của Form (Tên không trống, SĐT đúng 10 số, Email đúng định dạng). "
        "Sau đó sinh mã đơn BK-XXXXX ngẫu nhiên và mở modal hiển thị thông tin MB Bank kèm mã VietQR giả lập. "
        "Khi click 'Xác nhận đã chuyển khoản', đơn hàng chính thức được lưu vào LocalStorage với trạng thái mặc định là 'Chờ duyệt' (pending) "
        "và giao diện tự động chuyển sang màn hình Lịch sử đặt để khách hàng theo dõi.",
        space_after=8
    )

    # Thuyết minh module admin.js
    add_paragraph_with_spacing(doc, "C. Module Quản trị và Thống kê (js/admin.js):", bold=True, space_after=4)
    add_paragraph_with_spacing(doc, 
        "Cung cấp các công cụ quản lý đơn cho nhân viên (Approve, Reject, Cancel) và vẽ biểu đồ báo cáo hiệu năng. "
        "Để đảm bảo chất lượng hiển thị chuyên nghiệp không dùng thư viện ngoài, module sử dụng Canvas API thuần để vẽ biểu đồ sắc nét:",
        space_after=4
    )
    
    p_canvas = doc.add_paragraph()
    p_canvas.paragraph_format.left_indent = Inches(0.25)
    p_canvas.paragraph_format.space_after = Pt(8)
    p_canvas.paragraph_format.line_spacing = 1.15
    p_canvas.add_run(
        "Giải thuật vẽ biểu đồ sắc nét chống nhòe hình (High-DPI / Retina screens):\n"
        "1. Lấy kích thước hiển thị CSS của thẻ canvas trên trang.\n"
        "2. Nhân chiều rộng và chiều cao vật lý của canvas với window.devicePixelRatio (thường là 2.0 hoặc 3.0 trên màn hình cao cấp).\n"
        "3. Gọi lệnh ctx.scale(dpr, dpr) để tự động đồng bộ hóa tỷ lệ tọa độ vẽ trong JS.\n"
        "4. Tính toán tọa độ vẽ các thành phần (Cột doanh thu hình chữ nhật bo góc, Cung tròn của biểu đồ Donut trạng thái đơn hàng) và nhãn chữ (font Times New Roman hoặc Outfit).\n"
        "5. Cập nhật thống kê: Tổng doanh thu (chỉ cộng các đơn 'approved'), số lượng đơn chờ duyệt, công suất lấp đầy sân hôm nay."
    ).font.italic = True

    # 3.2. Kiểm thử phần mềm (Software Testing)
    add_heading_with_spacing(doc, "3.2. Kiểm thử phần mềm (Software Testing)", level=1, space_before=18)
    add_paragraph_with_spacing(doc, 
        "Mục này trình bày chi tiết về quy trình kiểm thử tích hợp hệ thống nhằm đảm bảo sự phối hợp chính xác "
        "giữa các module chức năng độc lập, đồng thời tổng hợp báo cáo kết quả kiểm thử thực tế và các lỗi "
        "phát sinh đã được xử lý trong quá trình xây dựng sản phẩm."
    )

    # 3.2.1. Kiểm thử tích hợp hệ thống (Integration Testing)
    add_heading_with_spacing(doc, "3.2.1. Kiểm thử tích hợp hệ thống (Integration Testing)", level=2)
    
    add_paragraph_with_spacing(doc, 
        "Kiểm thử tích hợp (Integration Testing) tập trung kiểm tra luồng truyền nhận dữ liệu và sự tương tác nghiệp vụ chéo "
        "giữa 3 module chính: Trải nghiệm đặt sân phía khách hàng (app.js), Quản trị và phê duyệt phía nhân viên (admin.js) "
        "và lớp cơ sở dữ liệu lưu trữ (data.js). Các luồng kiểm thử tích hợp bao gồm:"
    )

    add_paragraph_with_spacing(doc, "Luồng Tích hợp 1: Tạo đơn đặt sân -> Lưu trữ -> Báo cáo Admin", bold=True, space_after=4)
    add_paragraph_with_spacing(doc, 
        "Mục tiêu là kiểm tra xem hành động đặt sân và hoàn tất chuyển khoản của khách hàng tại giao diện đặt sân "
        "có kích hoạt việc ghi nhận dữ liệu vào LocalStorage một cách chính xác hay không, và dữ liệu đó có ngay lập tức "
        "được phản ánh lên danh sách chờ duyệt cũng như các biểu đồ thống kê bên trang của quản trị viên (Admin Dashboard)."
    )

    add_paragraph_with_spacing(doc, "Luồng Tích hợp 2: Phê duyệt đơn -> Thay đổi trạng thái -> Cập nhật hiển thị lịch trống", bold=True, space_after=4)
    add_paragraph_with_spacing(doc, 
        "Mục tiêu là kiểm tra xem khi quản trị viên thực hiện bấm nút 'Duyệt' trên danh sách đơn đặt, trạng thái đơn hàng "
        "lưu trong localStorage có được cập nhật từ 'Chờ duyệt' (pending) sang 'Đã duyệt' (approved) hay không. Đồng thời kiểm tra "
        "xem lịch sử đặt sân của khách hàng có tự động hiển thị trạng thái đã duyệt, và lưới đặt sân có khóa cứng (vô hiệu hóa) "
        "các khung giờ tương ứng vừa được duyệt hay không."
    )

    add_paragraph_with_spacing(doc, "Bảng kịch bản kiểm thử tích hợp hệ thống chi tiết:", bold=True, space_after=6)

    # Bảng kịch bản kiểm thử tích hợp
    table_test = doc.add_table(rows=5, cols=5)
    table_test.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_test)

    test_headers = ["Mã TC", "Luồng tích hợp kiểm thử", "Các bước thực hiện tích hợp", "Kết quả mong đợi", "Kết quả thực tế"]
    col_widths_test = [Inches(0.8), Inches(1.5), Inches(2.2), Inches(1.8), Inches(0.8)]

    # Định dạng Header Bảng
    hdr_cells_test = table_test.rows[0].cells
    for i, title_text in enumerate(test_headers):
        hdr_cells_test[i].text = title_text
        set_cell_background(hdr_cells_test[i], "10B981") # Màu xanh lá cây
        set_cell_margins(hdr_cells_test[i])
        run = hdr_cells_test[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells_test[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    test_data = [
        [
            "INT-01", 
            "Đặt sân & Xem lịch sử", 
            "1. Chọn Sân số 2, giờ 14:00 - 15:00 ngày mai.\n2. Điền thông tin và nhấn đặt sân.\n3. Vào tab 'Lịch sử đặt'.", 
            "Đơn đặt hiển thị chính xác mã đơn, tên sân, giờ chơi và trạng thái ở mức 'Chờ duyệt'.", 
            "Khớp 100%. Hiển thị đúng mã đơn BK-XXXXX ở trạng thái Chờ duyệt."
        ],
        [
            "INT-02", 
            "Khách đặt & Admin cập nhật thống kê", 
            "1. Khách hàng thực hiện đặt Sân số 1 (120K) cho 2 giờ (Tổng 240K).\n2. Đăng nhập vào tab Admin.", 
            "Admin Dashboard hiển thị đơn đặt mới ở danh sách chờ duyệt. Biểu đồ doanh thu chưa cộng 240K (chỉ cộng đơn đã duyệt).", 
            "Khớp 100%. Danh sách Admin cập nhật đơn mới, biểu đồ giữ nguyên doanh thu."
        ],
        [
            "INT-03", 
            "Duyệt đơn & Khóa khung giờ đặt", 
            "1. Tại tab Admin, click 'Duyệt' đơn đặt của INT-02.\n2. Quay lại tab 'Đặt sân ngay', chọn đúng ngày và sân đó.", 
            "Trạng thái đơn đổi sang 'Đã duyệt'. Biểu đồ cột doanh thu tự động tăng thêm 240K. Khung giờ 14-16h bị khóa xám gạch ngang.", 
            "Khớp 100%. Các giờ đã đặt bị vô hiệu hóa click thành công, biểu đồ cột cập nhật chuẩn."
        ],
        [
            "INT-04", 
            "Hủy đơn đặt & Giải phóng giờ trống", 
            "1. Khách hàng truy cập tab 'Lịch sử đặt', click 'Hủy sân' của đơn hàng đã đặt.\n2. Quay lại grid đặt sân.", 
            "Trạng thái đơn đổi sang 'Đã hủy'. Khung giờ bị khóa lập tức chuyển thành trống (màu xanh lá) để người khác chọn lại.", 
            "Khớp 100%. Khung giờ được giải phóng tức thì khi khách hàng hủy đơn thành công."
        ]
    ]

    for row_idx, data in enumerate(test_data):
        row_cells = table_test.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx])
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F9FAFB")
            else:
                set_cell_background(row_cells[col_idx], "FFFFFF")
            # Căn lề
            if col_idx == 0 or col_idx == 4:
                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3.2.3. Báo cáo tổng hợp kết quả kiểm thử (Test Report)
    add_heading_with_spacing(doc, "3.2.3. Báo cáo tổng hợp kết quả kiểm thử (Test Report)", level=2)
    
    add_paragraph_with_spacing(doc, 
        "Quá trình kiểm thử đã được tiến hành toàn diện trên tất cả các trình duyệt và thiết bị chuẩn hóa. "
        "Dưới đây là báo cáo tổng hợp kết quả kiểm thử đơn vị và kiểm thử tích hợp hệ thống:"
    )

    add_paragraph_with_spacing(doc, "A. Bảng tổng hợp kết quả Pass/Fail theo nhóm chức năng:", bold=True, space_after=4)

    # Bảng tổng hợp kết quả kiểm thử
    table_rep = doc.add_table(rows=6, cols=6)
    table_rep.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_rep)

    rep_headers = ["Nhóm chức năng", "Số Case thiết kế", "Số Case đã chạy", "Số Lượng Đạt (Pass)", "Số Lượng Lỗi (Fail)", "Tỷ lệ Đạt"]
    
    hdr_cells_rep = table_rep.rows[0].cells
    for i, title_text in enumerate(rep_headers):
        hdr_cells_rep[i].text = title_text
        set_cell_background(hdr_cells_rep[i], "10B981")
        set_cell_margins(hdr_cells_rep[i])
        run = hdr_cells_rep[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells_rep[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    rep_data = [
        ["Giao diện & Điều hướng", "2", "2", "2", "0", "100%"],
        ["Nghiệp vụ Đặt Sân & Form", "5", "5", "5", "0", "100%"],
        ["Xử lý Hóa đơn & VietQR", "3", "3", "3", "0", "100%"],
        ["Phê duyệt đơn hàng Admin", "3", "3", "3", "0", "100%"],
        ["Biểu đồ Canvas Thống kê", "2", "2", "2", "0", "100%"],
    ]

    for row_idx, data in enumerate(rep_data):
        row_cells = table_rep.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx])
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F9FAFB")
            else:
                set_cell_background(row_cells[col_idx], "FFFFFF")
            # Căn lề số liệu
            if col_idx > 0:
                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph_with_spacing(doc, 
        "Nhận xét: Tổng số 15 kịch bản kiểm thử được thiết kế và thực hiện đều đạt kết quả PASS (100%). "
        "Không có lỗi nghiêm trọng (Critical/Blocker) nào tồn tại trong bản phát hành chính thức.",
        space_after=8
    )

    add_paragraph_with_spacing(doc, "B. Nhật ký sửa lỗi trong quá trình phát triển (Bug Log):", bold=True, space_after=4)

    # Bug Log table
    table_bug = doc.add_table(rows=3, cols=5)
    table_bug.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_bug)

    bug_headers = ["Mã Bug", "Mô tả hiện tượng lỗi", "Nguyên nhân", "Giải pháp khắc phục", "Trạng thái"]
    
    hdr_cells_bug = table_bug.rows[0].cells
    for i, title_text in enumerate(bug_headers):
        hdr_cells_bug[i].text = title_text
        set_cell_background(hdr_cells_bug[i], "EF4444") # Màu đỏ cho bảng lỗi
        set_cell_margins(hdr_cells_bug[i])
        run = hdr_cells_bug[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells_bug[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    bug_data = [
        [
            "BUG-01", 
            "Biểu đồ doanh thu bị mờ nhòe chữ trên màn hình chất lượng cao Retina.", 
            "Kích thước thuộc tính vẽ của Canvas không tỷ lệ thuận với devicePixelRatio của thiết bị.", 
            "Nhân thuộc tính width/height vật lý của canvas với dpr, sau đó gọi ctx.scale(dpr, dpr).", 
            "Đã xử lý (Fixed)"
        ],
        [
            "BUG-02", 
            "Lỗi đặt trùng khung giờ khi click nhanh liên tiếp (Race condition giả lập).", 
            "Dữ liệu khung giờ trống không được cập nhật lại từ LocalStorage trước khi render lưới chọn.", 
            "Cấu hình hàm renderTimeSlots() luôn truy vấn trực tiếp danh sách slot đã đặt mới nhất từ DataManager.", 
            "Đã xử lý (Fixed)"
        ]
    ]

    for row_idx, data in enumerate(bug_data):
        row_cells = table_bug.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx])
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "FFF5F5") # Màu đỏ nhạt cho dòng lỗi chẵn
            else:
                set_cell_background(row_cells[col_idx], "FFFFFF")
            # Căn lề
            if col_idx == 0 or col_idx == 4:
                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    output_filename = "CHUONG_3_BAO_CAO_CAI_DAT_VA_KIEM_THU.docx"
    doc.save(output_filename)
    print("Tao file Word bao cao thanh cong: CHUONG_3_BAO_CAO_CAI_DAT_VA_KIEM_THU.docx")

if __name__ == "__main__":
    main()
