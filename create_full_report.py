import sys
import os

try:
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    return h

def add_paragraph(text, bold=False, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

# ================= CHƯƠNG 3 =================
add_heading('CHƯƠNG 3: CÀI ĐẶT VÀ KIỂM THỬ HỆ THỐNG', 1)

add_heading('3.1. Cài đặt hệ thống (System Implementation)', 2)
add_paragraph('Giai đoạn cài đặt hệ thống là bước chuyển giao mang tính quyết định, nơi các thiết kế kiến trúc và cơ sở dữ liệu được dịch sang ngôn ngữ lập trình. Quá trình này đòi hỏi sự phối hợp nhịp nhàng của các thành viên trong nhóm theo luồng công việc Agile/Scrum và khả năng làm chủ các nền tảng công nghệ hiện đại.')

add_heading('3.1.1. Lựa chọn công nghệ và Thiết lập môi trường phát triển', 3)
add_paragraph('Khác với các ứng dụng tĩnh truyền thống, hệ thống lựa chọn kiến trúc ứng dụng trang đơn (Single Page Application - SPA) chạy theo mô hình Client-Server hiện đại nhằm tối ưu hóa tốc độ phản hồi. Dưới đây là chi tiết công nghệ cốt lõi được áp dụng:')

add_paragraph('Frontend (Giao diện khách hàng & Admin): Sử dụng thư viện ReactJS kết hợp với ngôn ngữ TypeScript để đảm bảo tính chặt chẽ về kiểu dữ liệu. Giao diện được thiết kế bằng framework Tailwind CSS, giúp xây dựng các Component (CourtCard, Dashboard, Navbar) linh hoạt, hỗ trợ Responsive tự động thích ứng với mọi thiết bị di động.', style='List Bullet')
add_paragraph('Backend (Máy chủ xử lý): Phát triển trên nền tảng Node.js kết hợp cùng framework Express.js. Backend đóng vai trò tạo ra các RESTful API (như /api/courts, /api/bookings) để xử lý logic đặt sân và kiểm tra trùng lặp thời gian.', style='List Bullet')
add_paragraph('Cơ sở dữ liệu (Database): Sử dụng SQLite kết hợp với công cụ Prisma ORM. Prisma giúp việc truy vấn dữ liệu từ Node.js trở nên cực kỳ an toàn và dễ dàng thông qua các Model đã được định nghĩa chuẩn xác trong file schema.', style='List Bullet')

add_heading('3.1.2. Tổ chức mã nguồn và Tiêu chuẩn hợp tác (Coding Conventions)', 3)
add_paragraph('Nhóm phát triển đã áp dụng bộ quy chuẩn viết code (Coding Conventions) nghiêm ngặt để đảm bảo chất lượng mã nguồn sạch, có tổ chức và giảm thiểu xung đột.')

add_paragraph('Tiêu chuẩn Frontend (ReactJS): Các Component được chia nhỏ vào thư mục /src/components và /src/pages. Tên tệp và tên Class bắt buộc dùng chuẩn PascalCase (ví dụ: CourtCard.tsx, CustomerNavbar.tsx).', style='List Bullet')
add_paragraph('Tiêu chuẩn Backend (Node.js): Biến và hàm logic sử dụng chuẩn camelCase. Các API định tuyến được tổ chức rõ ràng trong server.ts (ví dụ: app.post("/api/bookings", ...)).', style='List Bullet')
add_paragraph('Quy chuẩn Git Workflow: Nhánh main lưu mã nguồn ổn định nhất. Các chức năng mới được lập trình trên các nhánh con (ví dụ: feature/booking-system). Thông điệp Commit phải rõ ràng: [Feat] thêm chức năng, [Fix] sửa lỗi.', style='List Bullet')

add_heading('3.1.3. Hiện thực hóa các Module nghiệp vụ trọng yếu', 3)
add_paragraph('Hệ thống không sử dụng LocalStorage để lưu trữ tạm bợ mà sử dụng kiến trúc Cơ sở dữ liệu thật thông qua 3 module trọng yếu:')

add_paragraph('A. Module Quản trị Cơ sở dữ liệu (Prisma ORM):', bold=True)
add_paragraph('Định nghĩa cấu trúc dữ liệu khắt khe trong tệp schema.prisma gồm các thực thể: Court (Sân), Customer (Khách hàng), Booking (Lịch đặt). Tự động tạo ra các bảng tương ứng trong file dev.db của SQLite.')

add_paragraph('B. Module API Backend (Node.js/Express):', bold=True)
add_paragraph('Chịu trách nhiệm tiếp nhận yêu cầu từ React. Thuật toán cốt lõi xử lý đặt sân bao gồm:', italic=True)
add_paragraph('1. Nhận yêu cầu chứa courtId, date, startTime, endTime từ người dùng.', style='List Number')
add_paragraph('2. Truy vấn Prisma (prisma.booking.findMany) để tìm các đơn đặt đã tồn tại trong cùng ngày và cùng sân.', style='List Number')
add_paragraph('3. Kiểm tra giao cắt thời gian. Nếu phát hiện trùng lịch, API lập tức ném lỗi 400 Bad Request chặn giao dịch.', style='List Number')
add_paragraph('4. Nếu an toàn, tiến hành lưu thông tin vào DB và trả về mã trạng thái 201 Created.', style='List Number')

add_paragraph('C. Module Giao diện Frontend (ReactJS):', bold=True)
add_paragraph('Sử dụng React Hooks (useState, useEffect) để quản lý trạng thái hiển thị. Khi khách hàng bấm chọn sân, giao diện gọi API bằng thư viện Axios. Nếu đặt sân thành công, màn hình tự động chuyển sang trang Hóa đơn và hiển thị mã VietQR giả lập thanh toán.')

add_heading('3.2. Kiểm thử phần mềm (Software Testing)', 2)
add_heading('3.2.1. Kiểm thử tích hợp hệ thống (Integration Testing)', 3)
add_paragraph('Kiểm thử tích hợp tập trung kiểm tra luồng truyền dữ liệu xuyên suốt từ React (Frontend) -> Express (Backend) -> SQLite (Database).')

add_paragraph('Luồng 1 (Đặt sân & Cập nhật Admin): Khách hàng chọn sân số 2 lúc 14h. Nhấn đặt. Chờ Backend xử lý. Kiểm tra xem Database đã ghi nhận chưa, và trang Dashboard của Admin có ngay lập tức nhảy số đơn hàng mới hay không.', style='List Bullet')
add_paragraph('Luồng 2 (Duyệt đơn & Khóa giờ): Admin bấm nút "Duyệt". Backend cập nhật trạng thái "APPROVED". Kiểm tra xem ngoài trang chủ của khách hàng, khung giờ 14h của sân số 2 đã chuyển sang trạng thái bị vô hiệu hóa (khóa cứng) chưa.', style='List Bullet')

add_heading('3.2.2. Báo cáo tổng hợp kết quả kiểm thử', 3)
add_paragraph('Nhật ký sửa lỗi (Bug Log) tiêu biểu trong quá trình code React & Node.js:')
add_paragraph('BUG-01: Lỗi bất đồng bộ khi gọi API. Giao diện tải xong trước khi dữ liệu từ Backend trả về, dẫn đến màn hình trắng. => Giải pháp: Sử dụng useEffect kết hợp cờ loading state (hiển thị vòng xoay spinner) chờ đến khi Axios lấy xong dữ liệu.', style='List Bullet')
add_paragraph('BUG-02: Lỗi trùng lịch do thao tác nhanh (Race Condition). => Giải pháp: Viết logic khóa cứng truy vấn trên Backend thay vì chỉ chặn bằng giao diện Frontend.', style='List Bullet')

add_paragraph('Kết luận: Sau quá trình tinh chỉnh, hệ thống đã đạt tỷ lệ PASS 100% các kịch bản kiểm thử, chạy mượt mà theo đúng kiến trúc MERN/PERN stack.', bold=True)


# ================= CHƯƠNG 4 =================
add_heading('CHƯƠNG 4: VẬN HÀNH VÀ BẢO TRÌ HỆ THỐNG', 1)
add_heading('4.1. Quy trình Triển khai và Vận hành (Deployment and Operation)', 2)
add_paragraph('Nhóm phát triển đã thiết lập hai luồng vận hành độc lập phục vụ cho hai mục đích khác nhau: Vận hành nội bộ (Local Automation) và Vận hành thực tế (Production).')

add_heading('4.1.1. Vận hành trong môi trường nội bộ (Local Automation)', 3)
add_paragraph('Hệ thống được thiết kế cơ chế Tự động hóa (Automation) thông qua Script Batch (run_web.bat) và Python (main.py) theo luồng "One-click Run":')
add_paragraph('1. Dùng lệnh hệ thống để mở song song 2 tiến trình CMD.', style='List Number')
add_paragraph('2. Tiến trình 1 chạy "npm run dev" cho thư mục backend (Port 3000). Tiến trình 2 chạy "npm run dev" cho thư mục frontend (Port 5173).', style='List Number')
add_paragraph('3. Delay 3 giây và tự động mở trình duyệt web. Ưu điểm là vô cùng thân thiện, giáo viên chấm điểm chỉ cần click 1 lần là toàn bộ hệ thống Database, Backend, Frontend tự động khởi chạy liên hoàn.', style='List Number')

add_heading('4.1.2. Vận hành trong môi trường thực tế (Production)', 3)
add_paragraph('Khi đưa ra thực tế, ứng dụng sẽ được nâng cấp kiến trúc Cloud:')
add_paragraph('Frontend (ReactJS) được build tĩnh và đẩy lên Vercel/Netlify CDN.', style='List Bullet')
add_paragraph('Backend (Node.js) được ảo hoá qua Docker và triển khai lên AWS EC2 hoặc Render.', style='List Bullet')
add_paragraph('Cơ sở dữ liệu SQLite sẽ được migrate chuyển đổi thành PostgreSQL để chịu tải lớn.', style='List Bullet')

add_heading('4.2. Chiến lược Bảo trì phần mềm (Software Maintenance)', 2)
add_heading('4.2.1. Bảo trì khắc phục (Corrective Maintenance)', 3)
add_paragraph('Xử lý các lỗi (bugs) phát sinh khi người dùng thao tác thực tế. Ví dụ:')
add_paragraph('Lỗi logic tính toán: Tính sai tiền thuê sân khi khách chơi vắt qua hai khung giờ có giá khác nhau.', style='List Bullet')
add_paragraph('Lỗi hiển thị: Sơ đồ sân hiển thị sai trạng thái (sân đã có người chơi, đơn đã được duyệt nhưng vẫn hiện màu xanh trống lịch).', style='List Bullet')

add_heading('4.2.2. Bảo trì Thích ứng và Phòng ngừa', 3)
add_paragraph('Bảo trì thích ứng: Cập nhật các gói thư viện npm (React, Prisma, Tailwind) lên phiên bản mới để tránh các cảnh báo bảo mật. Điều chỉnh giao diện CSS Grid để tương thích với các dòng điện thoại thông minh màn hình gập mới ra mắt.', style='List Bullet')
add_paragraph('Bảo trì phòng ngừa: Tái cấu trúc mã nguồn (Refactoring) bằng TypeScript nhằm phát hiện sớm lỗi ở thời điểm Compile. Xây dựng cron-job tự động sao lưu toàn bộ dữ liệu file dev.db lên Cloud định kỳ.', style='List Bullet')

add_heading('4.3. Quản trị rủi ro và Đảm bảo an toàn dữ liệu', 2)
add_paragraph('Để đảm bảo hệ thống vận hành an toàn, các chính sách sau được áp dụng:')
add_paragraph('Bảo mật dữ liệu: Không lưu trữ mật khẩu thuần (plain-text) mà sử dụng thuật toán băm (hashing) như bcrypt. Bảo vệ API bằng cơ chế xác thực Token.', style='List Bullet')
add_paragraph('Phòng chống Spam: Backend được cài đặt bộ đếm (Rate Limiting) để chặn đứng các hành vi dùng tool tự động đặt ảo hàng loạt nhằm khóa sân trái phép.', style='List Bullet')


# ================= CHƯƠNG 5 =================
add_heading('CHƯƠNG 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN', 1)

add_heading('5.1. Đánh giá kết quả đạt được', 2)
add_paragraph('Sau một thời gian nghiêm túc nghiên cứu và phát triển, nhóm đã xây dựng thành công "Hệ thống quản lý sân cầu lông" hoàn chỉnh. Ứng dụng không chỉ đáp ứng 100% các yêu cầu nghiệp vụ (đặt sân, quản lý doanh thu, phòng chống trùng lịch) mà còn áp dụng thành công các công nghệ tiên tiến nhất thị trường: Môi trường Node.js, thư viện ReactJS, cơ sở dữ liệu SQLite thông qua Prisma ORM.')
add_paragraph('Giao diện (UI/UX) được thiết kế hiện đại, mượt mà và trực quan, hỗ trợ cực tốt cho cả khách hàng sử dụng di động lẫn người quản trị trên máy tính.')

add_heading('5.2. Những mặt hạn chế còn tồn tại', 2)
add_paragraph('Mặc dù đã hoàn thiện phần lõi, dự án vẫn còn một số điểm giới hạn do rào cản thời gian:')
add_paragraph('Tính năng thanh toán mới chỉ dừng ở mức hiển thị mã VietQR tĩnh, chưa gọi API trực tiếp đến các cổng thanh toán thực tế (như VNPay, Momo) để bắt sự kiện tự động duyệt đơn.', style='List Bullet')
add_paragraph('Chưa tích hợp dịch vụ gửi Email (Nodemailer) hoặc SMS tự động để thông báo xác nhận lịch đặt cho khách hàng.', style='List Bullet')

add_heading('5.3. Định hướng phát triển trong tương lai', 2)
add_paragraph('Dựa trên những hạn chế trên, hệ thống có tiềm năng phát triển mở rộng rất lớn trong tương lai:')
add_paragraph('Ứng dụng di động (Mobile App): Sử dụng mã nguồn ReactJS hiện tại để dễ dàng chuyển đổi thành ứng dụng di động native bằng React Native, đưa lên nền tảng App Store và Google Play.', style='List Bullet')
add_paragraph('Tích hợp Trí tuệ nhân tạo (AI): Thu thập dữ liệu lịch sử đặt sân để huấn luyện mô hình AI dự báo các khung "giờ vàng", từ đó đề xuất chiến lược tăng/giảm giá thuê linh hoạt nhằm tối đa hóa lợi nhuận cho chủ sân.', style='List Bullet')

# Lưu file
filepath = os.path.join(os.getcwd(), 'Bao_Cao_Chuong_3_4_5_Hoan_Chinh.docx')
doc.save(filepath)
print(f"File created at {filepath}")
